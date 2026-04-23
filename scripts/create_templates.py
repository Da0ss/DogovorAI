"""
Script to generate DOCX templates for contract generation.
Run once to create the template files in app/templates_docx/.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "app" / "templates_docx"


def create_sale_contract_template():
    """Create sale contract (Договор купли-продажи) DOCX template."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ДОГОВОР КУПЛИ-ПРОДАЖИ")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # --- Date and number ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("г. Астана")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    date_para2 = doc.add_paragraph()
    date_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para2.add_run("Дата: {{date}}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # --- Preamble ---
    preamble = doc.add_paragraph()
    run = preamble.add_run(
        "{{seller_name}}, именуемый(ая) в дальнейшем «Продавец», с одной стороны, и "
        "{{buyer_name}}, именуемый(ая) в дальнейшем «Покупатель», с другой стороны, "
        "совместно именуемые «Стороны», заключили настоящий договор о нижеследующем:"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 1: Subject ---
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run("1. ПРЕДМЕТ ДОГОВОРА")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p1 = doc.add_paragraph()
    run = p1.add_run(
        "1.1. Продавец обязуется передать в собственность Покупателя, а Покупатель обязуется "
        "принять и оплатить товар (далее — «Товар») в соответствии с условиями настоящего Договора."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 2: Price ---
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run("2. ЦЕНА И ПОРЯДОК РАСЧЁТОВ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p2 = doc.add_paragraph()
    run = p2.add_run(
        "2.1. Общая стоимость Товара по настоящему Договору составляет {{amount}} тенге."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p3 = doc.add_paragraph()
    run = p3.add_run(
        "2.2. Оплата производится в безналичном порядке путём перечисления денежных средств "
        "на расчётный счёт Продавца в течение 5 (пяти) банковских дней с момента подписания "
        "настоящего Договора."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 3: Obligations ---
    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h3.add_run("3. ПРАВА И ОБЯЗАННОСТИ СТОРОН")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p4 = doc.add_paragraph()
    run = p4.add_run(
        "3.1. Продавец обязуется передать Товар надлежащего качества в сроки, "
        "установленные настоящим Договором."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p5 = doc.add_paragraph()
    run = p5.add_run(
        "3.2. Покупатель обязуется принять Товар и произвести оплату в порядке "
        "и в сроки, предусмотренные настоящим Договором."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 4: Liability ---
    h4 = doc.add_paragraph()
    h4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h4.add_run("4. ОТВЕТСТВЕННОСТЬ СТОРОН")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p6 = doc.add_paragraph()
    run = p6.add_run(
        "4.1. За неисполнение или ненадлежащее исполнение обязательств по настоящему "
        "Договору Стороны несут ответственность в соответствии с законодательством "
        "Республики Казахстан."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 5: Final ---
    h5 = doc.add_paragraph()
    h5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h5.add_run("5. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p7 = doc.add_paragraph()
    run = p7.add_run(
        "5.1. Настоящий Договор вступает в силу с момента подписания и действует "
        "до полного исполнения Сторонами своих обязательств."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p8 = doc.add_paragraph()
    run = p8.add_run(
        "5.2. Настоящий Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "по одному для каждой из Сторон."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # --- Signatures table ---
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    cell_seller = table.cell(0, 0)
    run = cell_seller.paragraphs[0].add_run("ПРОДАВЕЦ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    cell_buyer = table.cell(0, 1)
    run = cell_buyer.paragraphs[0].add_run("ПОКУПАТЕЛЬ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Signature row
    cell_seller_sign = table.cell(1, 0)
    run = cell_seller_sign.paragraphs[0].add_run("\n_________________\n{{seller_name}}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    cell_buyer_sign = table.cell(1, 1)
    run = cell_buyer_sign.paragraphs[0].add_run("\n_________________\n{{buyer_name}}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Save
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    output = TEMPLATES_DIR / "sale_contract.docx"
    doc.save(str(output))
    print(f"✅ Шаблон создан: {output}")


def create_labor_contract_template():
    """Create labor contract (Трудовой договор) DOCX template."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ТРУДОВОЙ ДОГОВОР")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # --- Date and city ---
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("г. Астана")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    date_para2 = doc.add_paragraph()
    date_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para2.add_run("Дата: {{start_date}}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # --- Preamble ---
    preamble = doc.add_paragraph()
    run = preamble.add_run(
        "{{employer_name}}, именуемый(ое) в дальнейшем «Работодатель», с одной стороны, и "
        "{{employee_name}}, именуемый(ая) в дальнейшем «Работник», с другой стороны, "
        "совместно именуемые «Стороны», заключили настоящий трудовой договор о нижеследующем:"
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 1: Subject ---
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run("1. ПРЕДМЕТ ДОГОВОРА")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p1 = doc.add_paragraph()
    run = p1.add_run(
        "1.1. Работодатель принимает Работника на должность «{{position}}» "
        "на условиях настоящего трудового договора."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p2 = doc.add_paragraph()
    run = p2.add_run(
        "1.2. Дата начала работы: {{start_date}}."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p3 = doc.add_paragraph()
    run = p3.add_run(
        "1.3. Настоящий трудовой договор заключён на неопределённый срок."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 2: Compensation ---
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run("2. ОПЛАТА ТРУДА")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p4 = doc.add_paragraph()
    run = p4.add_run(
        "2.1. Работнику устанавливается ежемесячная заработная плата в размере "
        "{{salary}} тенге."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p5 = doc.add_paragraph()
    run = p5.add_run(
        "2.2. Заработная плата выплачивается не позднее 10-го числа каждого месяца, "
        "следующего за расчётным."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 3: Work Schedule ---
    h3 = doc.add_paragraph()
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h3.add_run("3. РЕЖИМ РАБОТЫ И ОТДЫХА")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p6 = doc.add_paragraph()
    run = p6.add_run(
        "3.1. Работнику устанавливается 40-часовая рабочая неделя с двумя выходными днями "
        "(суббота и воскресенье)."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p7 = doc.add_paragraph()
    run = p7.add_run(
        "3.2. Работнику предоставляется ежегодный оплачиваемый трудовой отпуск "
        "продолжительностью 24 календарных дня в соответствии с Трудовым кодексом РК."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 4: Obligations ---
    h4 = doc.add_paragraph()
    h4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h4.add_run("4. ОБЯЗАННОСТИ СТОРОН")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p8 = doc.add_paragraph()
    run = p8.add_run(
        "4.1. Работодатель обязуется обеспечить Работнику условия труда, "
        "предусмотренные трудовым законодательством РК."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p9 = doc.add_paragraph()
    run = p9.add_run(
        "4.2. Работник обязуется добросовестно исполнять трудовые обязанности, "
        "соблюдать правила внутреннего трудового распорядка."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Section 5: Final ---
    h5 = doc.add_paragraph()
    h5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h5.add_run("5. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p10 = doc.add_paragraph()
    run = p10.add_run(
        "5.1. Настоящий договор составлен в двух экземплярах, имеющих одинаковую "
        "юридическую силу, по одному для каждой из Сторон."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    p11 = doc.add_paragraph()
    run = p11.add_run(
        "5.2. Споры, возникающие между Сторонами, разрешаются путём переговоров, "
        "а при недостижении согласия — в судебном порядке в соответствии с "
        "законодательством Республики Казахстан."
    )
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # --- Signatures table ---
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_employer = table.cell(0, 0)
    run = cell_employer.paragraphs[0].add_run("РАБОТОДАТЕЛЬ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    cell_employee = table.cell(0, 1)
    run = cell_employee.paragraphs[0].add_run("РАБОТНИК")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    cell_employer_sign = table.cell(1, 0)
    run = cell_employer_sign.paragraphs[0].add_run("\n_________________\n{{employer_name}}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    cell_employee_sign = table.cell(1, 1)
    run = cell_employee_sign.paragraphs[0].add_run("\n_________________\n{{employee_name}}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # Save
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    output = TEMPLATES_DIR / "labor_contract.docx"
    doc.save(str(output))
    print(f"✅ Шаблон создан: {output}")


if __name__ == "__main__":
    print("📝 Создание шаблонов DOCX...")
    create_sale_contract_template()
    create_labor_contract_template()
    print("✅ Все шаблоны созданы!")

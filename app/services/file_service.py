"""
File Service — Сервис для извлечения текста из документов различных форматов.
Поддерживает: PDF, DOCX, изображения (JPG/PNG) через OCR.
"""

import io
import os
import logging
import base64
from typing import Optional

from fastapi import UploadFile

from app.models.document import FileProcessResult, FileType
from config.settings import settings

logger = logging.getLogger(__name__)

# Vercel не имеет Tesseract в системе — OCR будет недоступен
# Проверяем наличие при загрузке и показываем предупреждение
_TESSERACT_AVAILABLE = False

# Пути к Tesseract OCR на разных системах (в порядке приоритета)
_TESSERACT_PATHS = [
    "/opt/homebrew/bin/tesseract",   # macOS Apple Silicon (M1/M2/M3)
    "/usr/local/bin/tesseract",      # macOS Intel
    "/usr/bin/tesseract",            # Linux
    "tesseract",                     # Системный PATH (fallback)
]


def _configure_tesseract() -> str:
    """
    Автоматически находит tesseract в системе и настраивает pytesseract.

    Returns:
        str: Путь к найденному tesseract или 'tesseract' если не найден
    """
    global _TESSERACT_AVAILABLE
    try:
        import pytesseract
        for path in _TESSERACT_PATHS:
            if path == "tesseract" or os.path.isfile(path):
                pytesseract.pytesseract.tesseract_cmd = path
                _TESSERACT_AVAILABLE = True
                logger.info(f"✅ Tesseract найден: {path}")
                return path
    except ImportError:
        pass
    logger.warning("⚠️ pytesseract не установлен или Tesseract не найден — OCR недоступен")
    _TESSERACT_AVAILABLE = False
    return "tesseract"


# Настраиваем Tesseract при загрузке модуля
_TESSERACT_CMD = _configure_tesseract()

_IS_SERVERLESS = (
    os.getenv("VERCEL") == "1"
    or os.getenv("AWS_LAMBDA_FUNCTION_NAME") is not None
)

# Максимальный размер файла: 20 МБ
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024

# Поддерживаемые MIME-типы
SUPPORTED_MIME_TYPES = {
    "application/pdf": FileType.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCX,
    "application/msword": FileType.DOCX,
    "image/jpeg": FileType.IMAGE,
    "image/jpg": FileType.IMAGE,
    "image/png": FileType.IMAGE,
    "image/tiff": FileType.IMAGE,
    "text/plain": FileType.TXT,
}


class FileService:
    """
    Сервис для обработки загруженных файлов.
    Извлекает текст из PDF, DOCX и изображений (через OCR).
    """

    @staticmethod
    def extract_text_from_image_google_vision(file_bytes: bytes) -> tuple[str, int]:
        """
        Extract text from an image with Google Cloud Vision REST API.

        Uses an API key instead of the Google SDK to keep Vercel dependencies
        small and serverless-friendly.
        """
        if not settings.google_cloud_vision_api_key:
            raise ValueError(
                "OCR изображений недоступен: GOOGLE_CLOUD_VISION_API_KEY не настроен."
            )

        try:
            import httpx

            image_content = base64.b64encode(file_bytes).decode("ascii")
            payload = {
                "requests": [
                    {
                        "image": {"content": image_content},
                        "features": [{"type": "TEXT_DETECTION"}],
                        "imageContext": {"languageHints": ["ru", "kk", "en"]},
                    }
                ]
            }
            url = (
                "https://vision.googleapis.com/v1/images:annotate"
                f"?key={settings.google_cloud_vision_api_key}"
            )
            with httpx.Client(timeout=30.0) as client:
                response = client.post(url, json=payload)
            response.raise_for_status()
            data = response.json()

            item = (data.get("responses") or [{}])[0]
            if item.get("error"):
                message = item["error"].get("message", "Google Vision OCR error")
                raise ValueError(f"Google Vision OCR error: {message}")

            annotations = item.get("textAnnotations") or []
            text = annotations[0].get("description", "").strip() if annotations else ""
            logger.info(f"✅ Google Vision OCR завершён: {len(text)} символов извлечено")
            return text, 1
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"❌ Google Vision OCR failed: {e}")
            raise ValueError(f"Ошибка OCR через Google Vision: {str(e)}")

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
        """
        Извлекает текст из PDF-файла через pypdf.
        Если извлечённый текст страницы пустой, пробует извлечь встроенные изображения
        и распознать текст на них через OCR.

        Args:
            file_bytes: Байты PDF-файла

        Returns:
            tuple: (извлечённый текст, количество страниц)

        Raises:
            ValueError: Если файл повреждён или не может быть обработан
        """
        try:
            from pypdf import PdfReader

            doc = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            page_count = len(doc.pages)

            for page_num in range(page_count):
                page = doc.pages[page_num]
                page_text = page.extract_text() or ""
                
                if page_text.strip():
                    text_parts.append(f"[Страница {page_num + 1}]\n{page_text}")
                else:
                    # Попробуем OCR картинок на этой странице (для отсканированных PDF)
                    page_images_text = []
                    for img_idx, img in enumerate(page.images):
                        try:
                            img_bytes = img.data
                            img_text, _ = FileService.extract_text_from_image_ocr(img_bytes)
                            if img_text.strip():
                                page_images_text.append(img_text)
                        except Exception as ocr_err:
                            logger.warning(
                                f"⚠️ Ошибка OCR картинки {img_idx} в PDF на стр {page_num + 1}: {ocr_err}"
                            )
                    
                    if page_images_text:
                        combined_text = "\n".join(page_images_text)
                        text_parts.append(f"[Страница {page_num + 1} (OCR)]\n{combined_text}")

            full_text = "\n\n".join(text_parts)
            logger.info(f"✅ PDF обработан: {page_count} стр., {len(full_text)} символов")
            return full_text, page_count

        except ImportError:
            logger.error("❌ pypdf не установлен. Выполните: pip install pypdf")
            raise ValueError("Не удалось обработать PDF")
        except Exception as e:
            logger.error(f"❌ Ошибка при обработке PDF: {str(e)}")
            raise ValueError(f"Не удалось обработать PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> tuple[str, int]:
        """
        Извлекает текст из DOCX-файла.

        Args:
            file_bytes: Байты DOCX-файла

        Returns:
            tuple: (извлечённый текст, количество параграфов)

        Raises:
            ValueError: Если файл повреждён или не является DOCX
        """
        try:
            # Проверяем сигнатуру файла (магические байты)
            if file_bytes.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
                raise ValueError(
                    "Формат .doc (Word 97-2003) не поддерживается. "
                    "Пожалуйста, сохраните файл в формате .docx и попробуйте снова."
                )
            if not file_bytes.startswith(b"PK\x03\x04"):
                raise ValueError("Файл не является корректным документом Word (.docx) или поврежден.")

            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n".join(paragraphs)
            para_count = len(paragraphs)
            logger.info(f"✅ DOCX обработан: {para_count} параграфов, {len(full_text)} символов")
            return full_text, para_count

        except ImportError:
            logger.error("❌ python-docx не установлен. Выполните: pip install python-docx")
            raise ValueError("Не удалось обработать DOCX")
        except Exception as e:
            error_msg = str(e)
            logger.error(f"❌ Ошибка при обработке DOCX: {error_msg}")
            if "is not a Word file" in error_msg or "themeManager" in error_msg:
                raise ValueError("Файл не является корректным документом Word (.docx) или поврежден.")
            if isinstance(e, ValueError):
                raise e
            raise ValueError(f"Не удалось обработать DOCX: {error_msg}")

    @staticmethod
    def extract_text_from_image_ocr(file_bytes: bytes) -> tuple[str, int]:
        """
        Извлекает текст из изображения с помощью Tesseract OCR.

        Args:
            file_bytes: Байты изображения (JPG/PNG/TIFF)

        Returns:
            tuple: (извлечённый через OCR текст, 1)

        Raises:
            ValueError: Если Tesseract не установлен или произошла ошибка OCR
        """
        if settings.google_cloud_vision_api_key:
            return FileService.extract_text_from_image_google_vision(file_bytes)

        if _IS_SERVERLESS or settings.is_production:
            raise ValueError(
                "Распознавание текста с картинок временно недоступно. "
                "Пожалуйста, загрузите документ в формате PDF, DOCX или TXT, либо обратитесь в поддержку."
            )

        try:
            from PIL import Image
            import pytesseract

            # Убеждаемся что путь к tesseract настроен
            _configure_tesseract()

            image = Image.open(io.BytesIO(file_bytes))

            # Применяем OCR с поддержкой русского и английского языков
            ocr_config = "--oem 3 --psm 3"

            # Пробуем сначала русский + английский, потом только английский
            available_langs = pytesseract.get_languages(config="")
            if "rus" in available_langs:
                lang = "rus+eng"
            else:
                lang = "eng"
                logger.warning(
                    "⚠️ Русский язык OCR не установлен. "
                    "Для полной поддержки: brew install tesseract-lang"
                )

            extracted_text = pytesseract.image_to_string(
                image,
                lang=lang,
                config=ocr_config
            )

            cleaned_text = extracted_text.strip()
            logger.info(f"✅ OCR завершён (lang={lang}): {len(cleaned_text)} символов извлечено")
            return cleaned_text, 1

        except ImportError as e:
            logger.error(f"❌ Зависимость для OCR не установлена: {str(e)}")
            raise ValueError(
                "OCR недоступен в данном окружении (Tesseract не установлен). "
                "Для локальной работы: brew install tesseract tesseract-lang && pip install pytesseract Pillow. "
                "На Vercel используйте PDF или DOCX вместо изображений."
            )
        except Exception as e:
            logger.error(f"❌ Ошибка OCR: {str(e)}")
            raise ValueError(f"Ошибка при распознавании текста: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> tuple[str, int]:
        """
        Извлекает текст из TXT-файла.

        Args:
            file_bytes: Байты TXT-файла

        Returns:
            tuple: (извлечённый текст, 1)

        Raises:
            ValueError: Если файл не может быть прочитан как текст
        """
        try:
            extracted_text = file_bytes.decode('utf-8')
            cleaned_text = extracted_text.strip()
            logger.info(f"✅ TXT обработан: {len(cleaned_text)} символов извлечено")
            return cleaned_text, 1
        except UnicodeDecodeError:
            try:
                # Fallback to general cyrillic
                extracted_text = file_bytes.decode('cp1251')
                cleaned_text = extracted_text.strip()
                logger.info(f"✅ TXT обработан (cp1251): {len(cleaned_text)} символов извлечено")
                return cleaned_text, 1
            except Exception as e:
                logger.error(f"❌ Ошибка декодирования TXT: {str(e)}")
                raise ValueError(f"Не удалось прочитать текстовый файл: {str(e)}")
        except Exception as e:
            logger.error(f"❌ Ошибка обработки TXT: {str(e)}")
            raise ValueError(f"Не удалось обработать TXT: {str(e)}")
    @classmethod
    async def process_uploaded_file(cls, file: UploadFile) -> FileProcessResult:
        """
        Основной метод: принимает загруженный файл, определяет тип и извлекает текст.

        Args:
            file: FastAPI UploadFile объект

        Returns:
            FileProcessResult: Результат обработки с извлечённым текстом
        """
        filename = file.filename or "unknown"
        content_type = file.content_type or ""

        logger.info(f"📄 Обработка файла: {filename} (тип: {content_type})")

        # Определяем тип файла
        file_type = SUPPORTED_MIME_TYPES.get(content_type.lower())

        # Если MIME-тип не определён — пробуем по расширению
        if file_type is None:
            ext = filename.lower().split(".")[-1] if "." in filename else ""
            ext_map = {
                "pdf": FileType.PDF,
                "docx": FileType.DOCX,
                "doc": FileType.DOCX,
                "jpg": FileType.IMAGE,
                "jpeg": FileType.IMAGE,
                "png": FileType.IMAGE,
                "txt": FileType.TXT,
            }
            file_type = ext_map.get(ext, FileType.UNKNOWN)

        if file_type == FileType.UNKNOWN:
            return FileProcessResult(
                filename=filename,
                file_type=FileType.UNKNOWN,
                success=False,
                error_message=(
                    f"Неподдерживаемый формат файла: {content_type}. "
                    "Поддерживаются: PDF, DOCX, TXT, JPG, PNG"
                )
            )

        # Читаем файл
        file_bytes = await file.read()

        # Проверка размера
        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            return FileProcessResult(
                filename=filename,
                file_type=file_type,
                success=False,
                error_message=f"Файл слишком большой: {len(file_bytes) // 1024 // 1024} МБ. Максимум: 20 МБ"
            )

        try:
            # Извлекаем текст в зависимости от типа
            if file_type == FileType.PDF:
                extracted_text, page_count = cls.extract_text_from_pdf(file_bytes)
            elif file_type == FileType.DOCX:
                extracted_text, page_count = cls.extract_text_from_docx(file_bytes)
            elif file_type == FileType.IMAGE:
                extracted_text, page_count = cls.extract_text_from_image_ocr(file_bytes)
            elif file_type == FileType.TXT:
                extracted_text, page_count = cls.extract_text_from_txt(file_bytes)
            else:
                raise ValueError("Неизвестный тип файла")

            if not extracted_text.strip():
                return FileProcessResult(
                    filename=filename,
                    file_type=file_type,
                    success=False,
                    error_message="Текст не найден в документе. Возможно, файл пустой или защищён."
                )

            return FileProcessResult(
                filename=filename,
                file_type=file_type,
                extracted_text=extracted_text,
                page_count=page_count,
                char_count=len(extracted_text),
                success=True
            )

        except ValueError as e:
            return FileProcessResult(
                filename=filename,
                file_type=file_type,
                success=False,
                error_message=str(e)
            )
        except Exception as e:
            logger.error(f"❌ Неожиданная ошибка при обработке {filename}: {str(e)}")
            return FileProcessResult(
                filename=filename,
                file_type=file_type,
                success=False,
                error_message=f"Внутренняя ошибка при обработке файла: {str(e)}"
            )


# Singleton-like экземпляр сервиса
file_service = FileService()

"""
DOCX Contract Generator Service.

Provides template-based document generation using python-docx.
Supports replacement of {{key}} placeholders in paragraphs and tables.

Architecture notes:
  - generate_docx() is the core engine (template-agnostic)
  - generate_sale_contract() / generate_labor_contract() are domain wrappers
  - Easy to extend with new contract types or add PDF conversion
"""

import logging
import re
import uuid
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Directories
BASE_DIR = Path(__file__).resolve().parent.parent.parent  # project root
TEMPLATES_DIR = BASE_DIR / "app" / "templates_docx"

# На Vercel файловая система read-only, кроме /tmp
# Используем /tmp/dogovorai_generated для serverless-совместимости
import os as _os
_is_vercel = _os.getenv("VERCEL") == "1" or _os.getenv("AWS_LAMBDA_FUNCTION_NAME") is not None
GENERATED_DIR = Path("/tmp/dogovorai_generated") if _is_vercel else BASE_DIR / "media" / "generated"


def _ensure_dirs() -> None:
    """Create output directories if they don't exist."""
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def _replace_placeholder(text: str, context: dict) -> str:
    """
    Replace all {{key}} placeholders in text with values from context.

    Args:
        text: Source text containing {{key}} placeholders
        context: Dictionary mapping keys to replacement values

    Returns:
        Text with all matched placeholders replaced
    """
    for key, value in context.items():
        placeholder = "{{" + key + "}}"
        if placeholder in text:
            text = text.replace(placeholder, str(value))
    return text


def _replace_in_paragraph(paragraph, context: dict) -> None:
    """
    Replace placeholders within a single paragraph.

    Handles the case where python-docx splits a placeholder across
    multiple runs by first checking the full paragraph text, then
    rebuilding runs if needed.

    Args:
        paragraph: python-docx Paragraph object
        context: Replacement context dictionary
    """
    full_text = paragraph.text
    if not any("{{" + k + "}}" in full_text for k in context):
        return

    # Simple case: placeholder is within a single run
    for run in paragraph.runs:
        if "{{" in run.text:
            run.text = _replace_placeholder(run.text, context)

    # If placeholders are split across runs, rebuild the paragraph
    remaining_text = paragraph.text
    if any("{{" + k + "}}" in remaining_text for k in context):
        new_text = _replace_placeholder(remaining_text, context)
        # Clear all runs except the first, set first run's text
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def generate_docx(template_path: str, context: dict, output_filename: Optional[str] = None) -> str:
    """
    Core DOCX generation engine.

    Loads a DOCX template, replaces all {{key}} placeholders in paragraphs
    and table cells, and saves the result to media/generated/.

    Args:
        template_path: Absolute or relative path to .docx template
        context: Dictionary of placeholder → value mappings
        output_filename: Optional custom filename (without extension).
                         If None, a UUID is generated.

    Returns:
        Absolute path to the generated .docx file

    Raises:
        FileNotFoundError: If the template file doesn't exist
        Exception: On DOCX processing errors
    """
    _ensure_dirs()

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Шаблон не найден: {template_path}")

    from docx import Document  # Lazy import: загружаем только при генерации

    doc = Document(str(template))

    # Replace in paragraphs (including headers/footers)
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, context)

    # Replace in headers and footers
    for section in doc.sections:
        for header_paragraph in section.header.paragraphs:
            _replace_in_paragraph(header_paragraph, context)
        for footer_paragraph in section.footer.paragraphs:
            _replace_in_paragraph(footer_paragraph, context)

    # Generate output filename
    if output_filename is None:
        output_filename = str(uuid.uuid4())
    output_path = GENERATED_DIR / f"{output_filename}.docx"

    doc.save(str(output_path))
    logger.info(f"📄 Документ сгенерирован: {output_path}")

    return str(output_path)


def generate_sale_contract(data: dict) -> str:
    """
    Generate a sale contract (Договор купли-продажи).

    Args:
        data: Dictionary with keys: seller_name, buyer_name, amount, date

    Returns:
        Path to the generated .docx file
    """
    template_path = str(TEMPLATES_DIR / "sale_contract.docx")
    context = {
        "seller_name": data["seller_name"],
        "buyer_name": data["buyer_name"],
        "amount": f"{data['amount']:,.2f}",
        "date": data["date"],
    }
    return generate_docx(template_path, context)


def generate_labor_contract(data: dict) -> str:
    """
    Generate a labor contract (Трудовой договор).

    Args:
        data: Dictionary with keys: employer_name, employee_name, salary, position, start_date

    Returns:
        Path to the generated .docx file
    """
    template_path = str(TEMPLATES_DIR / "labor_contract.docx")
    context = {
        "employer_name": data["employer_name"],
        "employee_name": data["employee_name"],
        "salary": f"{data['salary']:,.2f}",
        "position": data["position"],
        "start_date": data["start_date"],
    }
    return generate_docx(template_path, context)
